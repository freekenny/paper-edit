import streamlit as st
import google.generativeai as genai
from docx import Document
import io
import time
import hashlib
import difflib

# --- 页面配置 ---
st.set_page_config(layout="wide", page_title="Gemini 论文润色")
st.title("📑 Gemini 论文润色工坊")

# --- 侧边栏 ---
with st.sidebar:
    st.header("⚙️ 设置")
    
    # 1. 尝试读取系统配置的 Key
    try:
        system_key = st.secrets.get("GEMINI_API_KEY", "")
    except:
        system_key = ""

    # 2. 逻辑判断
    if system_key:
        # 如果后台配置了 Key
        api_key = system_key
        st.success("✅ 已激活公共 API Key")
        st.caption("你可以直接开始使用，无需输入密钥。")
        
        # (可选) 允许用户覆盖
        with st.expander("我想用自己的 Key"):
            user_input_key = st.text_input("覆盖默认 Key", type="password")
            if user_input_key:
                api_key = user_input_key
    else:
        # 如果后台没配置 Key，强制要求输入
        st.warning("⚠️ 未检测到公共 Key")
        api_key = st.text_input("请输入 Gemini API Key", type="password")
    
    st.divider()
    
    # 下面是原来的模型选择代码...
    model_choice = st.selectbox(
        "选择模型:",
        ["gemini-3-flash-preview", "gemini-2.5-flash", "gemini-2.0-flash"],
        index=0
    )
    
    style_option = st.selectbox("润色目标", 
        ("学术化润色", "修正语法", "重写/降重", "精简"))
    
    custom_instruction = st.text_area("额外指令", height=80)

# --- 核心函数 ---
def parse_docx(file):
    doc = Document(file)
    sections = []
    current_section = {"title": "文档开头", "content": ""}
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            if current_section["content"].strip():
                sections.append(current_section)
            current_section = {"title": para.text, "content": ""}
        else:
            current_section["content"] += para.text + "\n"
    if current_section["content"].strip():
        sections.append(current_section)
    return sections

def polish_text(text, style, custom_instr, api_key, model_name):
    """调用 Gemini API 润色文本"""
    try:
        genai.configure(api_key=api_key)
        
        model = genai.GenerativeModel(
            model_name=model_name,
            generation_config={
                "temperature": 1.0,
                "top_p": 0.95,
                "max_output_tokens": 8192,
            }
        )
        
        # 任务描述
        task_map = {
            "学术化润色": "transform this into formal academic writing with sophisticated vocabulary",
            "修正语法": "fix ALL grammar errors and improve sentence structure",
            "重写/降重": "completely rewrite this with different wording (change 70%+ of words)",
            "精简": "make this more concise by removing redundancy"
        }
        task = task_map.get(style, "improve this text significantly")
        
        prompt = f"""You MUST {task}.

Rules:
1. Fix all errors
2. Improve vocabulary and flow  
3. DO NOT return unchanged text
{f"4. {custom_instr}" if custom_instr else ""}

Original:
{text}

Improved:"""
        
        response = model.generate_content(prompt)
        
        if response and response.text:
            result = response.text.strip()
            # 去除可能的代码块标记
            if result.startswith("```"):
                result = "\n".join(result.split("\n")[1:-1]).strip()
            return result
        return None
    except Exception as e:
        st.error(f"API 错误: {str(e)}")
        return None

def compile_docx(data_list):
    doc = Document()
    for item in data_list:
        if item['title'] not in ["文档开头", "手动输入"]:
            doc.add_heading(item['title'], level=1)
        content = item['polished'] if item['polished'] else item['original']
        for line in content.split('\n'):
            if line.strip():
                doc.add_paragraph(line)
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def generate_diff_html(original, polished):
    """生成类似 Word 修订模式的 HTML 差异显示"""
    if not polished:
        return "<p style='color: gray; font-style: italic; font-size: 11px;'>暂无修改</p>"
    
    # 使用 difflib 计算差异
    diff = difflib.ndiff(original.split(), polished.split())
    
    html_parts = []
    html_parts.append("<div style='font-family: Arial, sans-serif; font-size: 11px; line-height: 1.6; padding: 8px; background: #ffffff; border-radius: 5px; border: 1px solid #e0e0e0; height: 200px; overflow-y: auto;'>")
    
    for item in diff:
        if item.startswith('  '):  # 未改变
            word = item[2:]
            html_parts.append(f"<span style='background-color: #686868; padding: 1px 2px; border-radius: 2px; font-size: 10px;'>{word}</span> ")
        elif item.startswith('- '):  # 删除
            word = item[2:]
            html_parts.append(f"<span style='background-color: #ffcccc; text-decoration: line-through; color: #cc0000; padding: 1px 2px; border-radius: 2px; font-size: 10px;'>{word}</span> ")
        elif item.startswith('+ '):  # 添加
            word = item[2:]
            html_parts.append(f"<span style='background-color: #ccffcc; color: #006600; font-weight: bold; padding: 1px 2px; border-radius: 2px; font-size: 10px;'>{word}</span> ")
    
    html_parts.append("</div>")
    html_parts.append("<div style='margin-top: 5px; font-size: 9px; color: #666;'>")
    html_parts.append("<span style='background-color: #686868; color: #ffffff; padding: 1px 4px; border-radius: 2px; margin-right: 3px;'>未改变</span>")
    html_parts.append("<span style='background-color: #ffcccc; padding: 1px 4px; border-radius: 2px; margin-right: 3px;'>删除</span>")
    html_parts.append("<span style='background-color: #ccffcc; padding: 1px 4px; border-radius: 2px;'>添加</span>")
    html_parts.append("</div>")
    
    return ''.join(html_parts)

# --- Session State 初始化（使用新的变量名） ---
if "document_data" not in st.session_state:
    st.session_state.document_data = []

if "widget_counter" not in st.session_state:
    st.session_state.widget_counter = 0

# --- 文件上传 ---
uploaded_file = st.file_uploader("上传 .docx 文件", type=['docx'])

col1, col2 = st.columns([3, 1])
with col1:
    manual_input = st.text_area("或粘贴文本", height=100)
with col2:
    st.write("\n")
    if st.button("📝 载入文本", use_container_width=True):
        if manual_input.strip():
            st.session_state.document_data = [{
                "id": hashlib.md5(manual_input.encode()).hexdigest()[:8],
                "title": "手动输入",
                "original": manual_input,
                "polished": ""
            }]
            st.success("✅ 已载入")
            st.rerun()

if uploaded_file:
    file_id = hashlib.md5((uploaded_file.name + str(uploaded_file.size)).encode()).hexdigest()[:8]
    if not st.session_state.document_data or st.session_state.document_data[0].get('file_id') != file_id:
        sections = parse_docx(uploaded_file)
        st.session_state.document_data = []
        for sec in sections:
            st.session_state.document_data.append({
                "id": hashlib.md5(sec['content'].encode()).hexdigest()[:8],
                "title": sec['title'],
                "original": sec['content'],
                "polished": "",
                "file_id": file_id
            })
        st.success(f"✅ 载入 {len(sections)} 段")
        st.rerun()

st.divider()

# --- 显示和处理 ---
if len(st.session_state.document_data) > 0:
    
    col_a, col_b = st.columns([3, 1])
    col_a.info(f"共 {len(st.session_state.document_data)} 段")
    if col_b.button("🗑️ 清空", use_container_width=True):
        st.session_state.document_data = []
        st.rerun()
    
    for idx, item in enumerate(st.session_state.document_data):
        status = "✅ 已润色" if item['polished'] else "⏳ 待处理"
        
        with st.expander(f"{item['title']} [{status}]", expanded=True):
            col1, col2, col3, col4 = st.columns([3.5, 3, 0.5, 3.5])
            
            with col1:
                st.caption("📄 原文")
                st.text_area(
                    "原文", 
                    item['original'], 
                    height=200, 
                    key=f"orig_{item['id']}_{st.session_state.widget_counter}", 
                    disabled=True, 
                    label_visibility="collapsed"
                )
            
            with col2:
                st.caption("📝 修订追踪")
                # 显示差异
                if item['polished']:
                    diff_html = generate_diff_html(item['original'], item['polished'])
                    st.markdown(diff_html, unsafe_allow_html=True)
                else:
                    st.markdown("<div style='height: 200px; display: flex; align-items: center; justify-content: center; background: #f9f9f9; border-radius: 5px; border: 1px solid #e0e0e0;'><p style='color: gray; font-style: italic; font-size: 12px;'>点击润色按钮查看修改</p></div>", unsafe_allow_html=True)
            
            with col3:
                st.write("\n\n")
                if st.button("⚡", key=f"btn_{item['id']}", use_container_width=True, help="润色"):
                    if not api_key:
                        st.error("请输入 API Key")
                    else:
                        with st.spinner("处理中..."):
                            result = polish_text(
                                item['original'],
                                style_option,
                                custom_instruction,
                                api_key,
                                model_choice
                            )
                            
                            if result and result != item['original']:
                                # 更新数据
                                st.session_state.document_data[idx]['polished'] = result
                                # 增加计数器，强制刷新所有 widget
                                st.session_state.widget_counter += 1
                                st.success("✅ 完成")
                                time.sleep(0.5)
                                st.rerun()
                            elif result == item['original']:
                                st.warning("⚠️ 结果与原文相同")
                            else:
                                st.error("❌ 失败")
            
            with col4:
                st.caption("✨ 润色结果")
                display_text = item['polished'] if item['polished'] else item['original']
                
                # 使用带计数器的唯一 key
                new_text = st.text_area(
                    "结果",
                    display_text,
                    height=200,
                    key=f"result_{item['id']}_{st.session_state.widget_counter}",
                    label_visibility="collapsed"
                )
                
                # 检测手动编辑
                if new_text != display_text:
                    if st.button("💾 保存编辑", key=f"save_{item['id']}"):
                        st.session_state.document_data[idx]['polished'] = new_text
                        st.session_state.widget_counter += 1
                        st.success("已保存")
                        st.rerun()
    
    st.divider()
    
    if st.button("📥 导出 Word", type="primary"):
        data = compile_docx(st.session_state.document_data)
        st.download_button(
            "⬇️ 下载",
            data,
            "polished.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
else:
    st.info("👋 请上传文件或粘贴文本")